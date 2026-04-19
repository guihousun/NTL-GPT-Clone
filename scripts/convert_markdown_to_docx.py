"""Convert a Markdown manuscript to a formatted DOCX document."""

from __future__ import annotations

import argparse
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_east_asian_font(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.color.rgb = RGBColor(31, 78, 121)


def add_runs_from_node(paragraph, node: Tag | NavigableString) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            set_east_asian_font(run, "SimSun")
        return

    if node.name == "br":
        paragraph.add_run().add_break(WD_BREAK.LINE)
        return

    text = node.get_text()
    if not text:
        return
    run = paragraph.add_run(text)
    set_east_asian_font(run, "SimSun")
    if node.name in {"strong", "b"}:
        run.bold = True
    if node.name in {"em", "i"}:
        run.italic = True
    if node.name == "code":
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_paragraph_from_tag(document: Document, tag: Tag, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    for child in tag.children:
        add_runs_from_node(paragraph, child)


def add_table_from_tag(document: Document, tag: Tag) -> None:
    rows = tag.find_all("tr")
    if not rows:
        return
    max_cols = max(len(row.find_all(["th", "td"])) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            target = table.cell(row_idx, col_idx)
            target.text = cell.get_text(" ", strip=True)
            for paragraph in target.paragraphs:
                for run in paragraph.runs:
                    set_east_asian_font(run, "SimSun")
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    document.add_paragraph()


def add_list_from_tag(document: Document, tag: Tag, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in tag.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        for child in li.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                continue
            add_runs_from_node(paragraph, child)


def convert_markdown(input_path: Path, output_path: Path) -> None:
    text = input_path.read_text(encoding="utf-8")
    html = markdown.markdown(text, extensions=["tables", "fenced_code", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")

    document = Document()
    style_document(document)

    for node in soup.contents:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue
        if node.name == "h1":
            document.add_paragraph(node.get_text(" ", strip=True), style="Title")
        elif node.name == "h2":
            document.add_paragraph(node.get_text(" ", strip=True), style="Heading 1")
        elif node.name == "h3":
            document.add_paragraph(node.get_text(" ", strip=True), style="Heading 2")
        elif node.name == "h4":
            document.add_paragraph(node.get_text(" ", strip=True), style="Heading 3")
        elif node.name == "p":
            add_paragraph_from_tag(document, node)
        elif node.name == "table":
            add_table_from_tag(document, node)
        elif node.name == "ul":
            add_list_from_tag(document, node, ordered=False)
        elif node.name == "ol":
            add_list_from_tag(document, node, ordered=True)
        elif node.name == "pre":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(node.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif node.name == "hr":
            document.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", help="Input Markdown path.")
    parser.add_argument("output", help="Output DOCX path.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    convert_markdown(Path(args.input).resolve(), Path(args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
